from llama_cpp import Llama
import re

# import translators as ts
# from googletrans import Translator

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


import argostranslate.package
import argostranslate.translate




def translate(text):
    # translated = None
    # while translated is None:
    #     try:
    #         translated = PyDeepLX.translate(text, "US", "RU") # Return String
    #     except Exception as e:
    #         print(f'{str(e)}')
    # return translated

    # text = f'''Translate into russian: {text}\nПеревод: '''
    # count_of_characters = len(text)
    # output = llm(f"{text}", temperature=0.5, echo=True, max_tokens=-1)
    # translated = output["choices"][0]["text"]
    # return translated[count_of_characters:]

    # translator = Translator()
    # translated = translator.translate(text, src='en', dest='ru').text

    # return translated

    from_code = "en"
    to_code = "ru"

    # Download and install Argos Translate package
    argostranslate.package.update_package_index()
    available_packages = argostranslate.package.get_available_packages()
    package_to_install = next(
        filter(
            lambda x: x.from_code == from_code and x.to_code == to_code, available_packages
        )
    )
    argostranslate.package.install_from_path(package_to_install.download())

    # Translate
    translatedText = argostranslate.translate.translate(text, from_code, to_code)
    return translatedText





def create_signature(doc):

    director_signature = doc.add_paragraph()
    director_signature.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    text = f'APPROVED\n Director of the Federal State Budgetary Institution\n "Administration of {project}"\nBorisenko I.A.\n____________________\n"July 12", 2022\n'
    director_signature.add_run(translate(text))

    for _ in range(7):  # Adjust the range for more or less space
        doc.add_paragraph()



def create_title(doc, title):
    heading = doc.add_paragraph()
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = heading.add_run(f'{translate(title)}\n')
    run.bold = True
    

def create_description(doc, description):
    heading = doc.add_paragraph()
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    heading.add_run(f'{translate(description)}')


def decide_description(llm, title):
    try: 
        text = f'''Q: This document is essential for the planning, operation, and management of a business. Title of Business Documentation for the power plant operation, which is called {project}. Title is "{title}". Give little description of this document. No more than 10 words \nA: Sure, here's a concise description of the document: "'''
        output = llm(f"{text}", temperature=2, echo=True, max_tokens=-1)
        text = output["choices"][0]["text"]

        description = re.findall(r'\"(.*?)\"', text)

        return description[-1]
        


    except Exception as e:
        print(f'{str(e)}')
        return create_title(llm, doc, title)


def decide_title(llm, project):
    try:
        text = f'''Q: Here's documents are essential for the planning, operation, and management of a business: Design and Engineering Documents, Environmental Permits, Safety and Emergency Plans, Maintenance and Operations Manuals, Compliance Certificates, Financial Agreement, Insurance Policies, Workers' Credentials and Training Records, Utility Interconnection Agreements, Marketing and Sales Agreements
        Give me 5 titles from Business Documentation for the power plant operation, which is called {project}. No more than 7 words \nA: Well here's the name of the document you need, which if leaked will get your business in trouble:\n 1. "'''
        output = llm(f"{text}", temperature=2, echo=True, max_tokens=-1)

        text = output["choices"][0]["text"]
        phrases = re.findall(r'\"(.*?)\"', text)

        if phrases[4] is None:
            raise ValueError("Не сгенерировано название документа")
        return phrases[4]
    
    except Exception as e:
        print(f'{str(e)}')
        return decide_title(llm)

def decide_list_of_contents(llm, title, description, project):
    # try:
        text = f'''Title: {title}
Description: {description}
Note to List of Contents: Make less than 8 points

List of Contents:
1. General Provisions
2.'''
        output = llm(f"{text}", temperature=2, echo=True, max_tokens=-1)

        text = output["choices"][0]["text"]
        print(text)

        contents = re.findall(r'\d+\.\s(.+)', text)
        print(contents)
        return contents[:7]
        



def create_contents(doc, llm, title, description, list_of_contents):
    z = 1
    for content in list_of_contents:
        # content = list_of_contents[0]
        # print(content)
        heading = doc.add_paragraph()
        # heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        heading.add_run(f'{translate(f"{z}. {content}")}')
        heading.bold = True

        text_content = doc.add_paragraph()
                
        text = f'''Title: {title}\nDescription: {description}\nContent: {content}\n Note to text: The text should have a maximum of 1000 words. Indent the first line\nText: '''

        output = llm(f"{text}", temperature=2, echo=True, max_tokens=-1)

        text = output["choices"][0]["text"]
        print(text)
        text_content.add_run(translate(text))
        # text_content.add_run(text)
        z+=1


    
    # except Exception as e:
    #     print(f'{str(e)}')
    #     return decide_title(llm)

llm = Llama(model_path="C:/Users/user/Downloads/WizardLM-2-7B.Q5_K_M.gguf", n_threads=16, n_gpu_layers=33, n_ctx=8192)
project = "Novovoronezh Nuclear Power Plant"
title = decide_title(llm, project)
description = decide_description(llm, title)

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)


# title = "Radiation Emergency Preparedness and Response"
# description = "Designed for use by officials of the Novovoronezh NPP, emergency response teams, and other individuals responsible for radiation protection and emergency response during a radiological incident."


create_signature(doc)
create_title(doc, title)


create_description(doc, description)
doc.add_page_break()
list_of_contents = decide_list_of_contents(llm, title, description, project)
create_contents(doc, llm, title, description, list_of_contents)



doc.save(f'words/{title}.docx')



# 
# print(translated)
